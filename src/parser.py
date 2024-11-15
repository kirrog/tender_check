import argparse
import json
import subprocess
from pprint import pprint

import PyPDF2
import docx
import pandas as pd
import requests
from tabula import read_pdf


def parse_data_from_url(url_text: str):
    resp = requests.get(f'https://zakupki.mos.ru/newapi/api/Auction/Get?auctionId={url_text.rsplit("/", 1)[-1]}')
    resp_json = resp.json()

    for item in resp_json['items']:
        item_resp = requests.get(
            f'https://zakupki.mos.ru/newapi/api/Auction/GetAuctionItemAdditionalInfo?itemId={item["id"]}')
        item['specification'] = item_resp.json()

    for file in resp_json['files']:
        file_name = file["name"]
        url = f'https://zakupki.mos.ru/newapi/api/FileStorage/Download?id={file["id"]}'
        data = b''

        r = requests.get(url)
        for chunk in r.iter_content(chunk_size=512 * 1024):
            if chunk:
                data += chunk
        file_type = ""
        pdf_file_path = "../tmp/tmp.pdf"
        docx_file_path = "../tmp/tmp.docx"
        doc_file_path = "../tmp/tmp.doc"
        if ".pdf" in file_name:
            file_type = "pdf"
            with open(pdf_file_path, "wb") as f:
                f.write(data)
        elif ".docx" in file_name:
            file_type = "docx"
            with open(docx_file_path, "wb") as f:
                f.write(data)
        elif ".doc" in file_name:
            file_type = "doc"
            with open(doc_file_path, "wb") as f:
                f.write(data)

        if len(file_type) == 0:
            print(f"Wrong endfix!!!: {file_name}")
            file["text"] = ""
            file["type"] = "err"
            continue

        text = ""
        tables = []
        if file_type == "pdf":
            reader = PyPDF2.PdfReader(pdf_file_path)
            tables_raw = read_pdf(pdf_file_path, pages="all")

            pdf_texts = []
            for i in range(len(reader.pages)):
                page_text = reader.pages[i].extract_text()
                pdf_texts.append(page_text)

            for table in tables_raw:
                tables.append(table)
            text = "\n".join([" ".join(x.split()) for x in pdf_texts if x != ""])

        if file_type == "doc":
            output = subprocess.check_output(["soffice", "--headless", "--invisible", "--convert-to", "docx",
                                              doc_file_path, "--outdir", "../tmp/"])
            file_type = "docx"

        if file_type == "docx":
            doc = docx.Document(str(docx_file_path))
            texts = []
            for paragraph in doc.paragraphs:
                texts.append(paragraph.text)
            for table in doc.tables:
                column_list = []
                for col in table.columns:
                    l = []
                    column_list.append(l)
                    for cell in col.cells:
                        l.append(str(cell.text))
                tables.append(pd.DataFrame(column_list).T)
            text = "\n".join([" ".join(x.split()) for x in texts if x != ""])

        file['text'] = text + "\n" + "\n".join([json.dumps(x.to_dict('dict'), ensure_ascii=False) for x in tables])
        file['tables'] = [x.to_dict('dict') for x in tables]
        file['type'] = file_type
    return resp_json


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Parse KC')
    parser.add_argument('-url', '--url', help='URL for KC', default="https://zakupki.mos.ru/auction/9864533")
    args = parser.parse_args()
    resp_json = parse_data_from_url(args.url)
    with open("../cases/9864533/reponse.json", "w", encoding="utf-8") as f:
        json.dump(resp_json, f, ensure_ascii=False)
    pprint(resp_json)
